"""
Extract 3P Programs DOCX into structured JSON data.
Outputs: src/utils/programs/data/3p-programs.json
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


SRC = Path(r"C:\Users\nikjo\OneDrive\Desktop\EverWatt_Engine\3P_PROGRAMS\California Energy Efficiency Programs Research.docx")
DEST = Path("src/utils/programs/data/3p-programs.json")

TRANS = {
    "\u2019": "'",
    "\u2018": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2013": "-",
    "\u2014": "-",
    "\u2026": "...",
    "\u2022": "*",
    "\ufeff": "",
    "\xa0": " ",
    "\u2265": ">=",
    "\u2264": "<=",
}


def normalize_text(text: str) -> str:
    """Normalize text by replacing special characters."""
    for k, v in TRANS.items():
        text = text.replace(k, v)
    return text.encode("ascii", "ignore").decode("ascii").strip()


def extract_programs(doc: Document) -> list[dict]:
    """Extract program information from the document."""
    programs = []
    current_program = None
    current_section = None
    
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue
        
        # Detect program names (usually bold or numbered sections)
        # Look for patterns like "2.1 Program Name" or standalone program names
        if re.match(r'^\d+\.\d+\s+[A-Z]', text) or (len(text) > 10 and len(text) < 100 and text[0].isupper()):
            # Save previous program if exists
            if current_program:
                programs.append(current_program)
            
            # Start new program
            current_program = {
                "name": text,
                "section": current_section,
                "description": "",
                "details": [],
                "utility": None,
                "implementer": None,
                "category": None,
            }
        elif current_program:
            # Check for utility mentions
            utilities = ["PG&E", "SDG&E", "SCE", "SoCalGas", "SCG", "LADWP", "SMUD"]
            for util in utilities:
                if util in text:
                    current_program["utility"] = util
                    break
            
            # Check for implementer mentions
            implementers = ["CLEAResult", "Energy Solutions", "TRC", "Willdan"]
            for impl in implementers:
                if impl in text:
                    current_program["implementer"] = impl
                    break
            
            # Add to description or details
            if "Lead Administrator:" in text or "Implementer:" in text:
                current_program["details"].append(text)
            elif len(text) > 50:
                if not current_program["description"]:
                    current_program["description"] = text
                else:
                    current_program["details"].append(text)
    
    # Add last program
    if current_program:
        programs.append(current_program)
    
    return programs


def main() -> None:
    doc = Document(SRC)
    
    # Extract full text for now (we can structure it better later)
    paragraphs = []
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if text:
            paragraphs.append(text)
    
    # Try to extract structured programs
    programs = extract_programs(doc)
    
    output = {
        "source_file": str(SRC),
        "total_paragraphs": len(paragraphs),
        "programs": programs,
        "full_text": "\n\n".join(paragraphs),
    }
    
    DEST.parent.mkdir(parents=True, exist_ok=True)
    with DEST.open("w", encoding="utf-8") as f:
        json.dump(output, f, indent=2)
    
    print(f"Wrote {DEST} with {len(programs)} programs extracted from {len(paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
